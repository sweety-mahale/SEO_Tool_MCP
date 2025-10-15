"""
Google Drive MCP Integration for SEO Analyzer
Saves SEO analysis results as Word documents to Google Drive

"""

import os
from typing import Dict, Optional
from datetime import datetime
from google.oauth2.credentials import Credentials
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import tempfile
from pathlib import Path

# Scopes required for Google Drive
SCOPES = ['https://www.googleapis.com/auth/drive.file']

class GoogleDriveSEOExporter:
    """Export SEO analysis to Google Drive as Word document"""
    
    def __init__(self, credentials_path: str = "credentials.json", token_path: str = "token.json"):
        """
        Initialize Google Drive exporter.
        
        Args:
            credentials_path: Path to OAuth2 credentials.json from Google Cloud Console
            token_path: Path to store authentication token
        """
        self.credentials_path = credentials_path
        self.token_path = token_path
        self.service = None
        self._authenticate()
    
    def _authenticate(self):
        """Authenticate with Google Drive API"""
        creds = None
        
        # Check if token.json exists (stored credentials)
        if os.path.exists(self.token_path):
            creds = Credentials.from_authorized_user_file(self.token_path, SCOPES)
        
        # If no valid credentials, let user log in
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                if not os.path.exists(self.credentials_path):
                    raise FileNotFoundError(
                        f"Credentials file not found at {self.credentials_path}. "
                        "Please download OAuth2 credentials from Google Cloud Console."
                    )
                
                flow = InstalledAppFlow.from_client_secrets_file(
                    self.credentials_path, SCOPES
                )
                creds = flow.run_local_server(port=0)
            
            # Save credentials for next run
            with open(self.token_path, 'w') as token:
                token.write(creds.to_json())
        
        self.service = build('drive', 'v3', credentials=creds)
    
    def _create_word_document(self, seo_data: Dict, output_path: str) -> str:
        """
        Create a Word document from SEO analysis data.
        
        Args:
            seo_data: SEO analysis results dictionary
            output_path: Local path to save the Word document
            
        Returns:
            Path to created document
        """
        doc = Document()
        
        # Extract data
        url = seo_data.get('url', 'N/A')
        primary_keyword = seo_data.get('primary_keyword', 'N/A')
        business_goals = seo_data.get('business_goals', {})
        metrics = seo_data.get('metrics', {})
        ai_rec = seo_data.get('ai_recommendations', {})
        issues = seo_data.get('issues_found', [])
        suggestions = seo_data.get('suggestions', [])
        
        # Title
        title = doc.add_heading('SEO Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"URL Analyzed: {url}")
        if primary_keyword:
            doc.add_paragraph(f"Primary Keyword: {primary_keyword}")
        
        doc.add_paragraph()  # Blank line
        
        # Executive Summary
        doc.add_heading('📊 Executive Summary', 1)
        overall_score = ai_rec.get('seo_score', 0)
        p = doc.add_paragraph()
        run = p.add_run(f"Overall SEO Score: {overall_score}/100")
        run.bold = True
        run.font.size = Pt(14)
        
        if ai_rec.get('score_explanation'):
            doc.add_paragraph(ai_rec['score_explanation'])
        
        # Score Breakdown
        score_breakdown = ai_rec.get('score_breakdown', {})
        if score_breakdown:
            doc.add_heading('Score Breakdown', 2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Category'
            hdr_cells[1].text = 'Score'
            
            for category, score in score_breakdown.items():
                row_cells = table.add_row().cells
                row_cells[0].text = category.replace('_', ' ').title()
                row_cells[1].text = f"{score}/100"
        
        doc.add_page_break()
        
        # Business Goals Alignment
        business_goal_alignment = ai_rec.get('business_goal_alignment', {})
        if business_goal_alignment:
            doc.add_heading('🎯 Business Goal Alignment', 1)
            
            alignment_score = business_goal_alignment.get('alignment_score', 0)
            p = doc.add_paragraph()
            run = p.add_run(f"Alignment Score: {alignment_score}/100")
            run.bold = True
            run.font.size = Pt(12)
            
            if business_goal_alignment.get('strengths'):
                doc.add_heading('Strengths', 2)
                for strength in business_goal_alignment['strengths']:
                    doc.add_paragraph(strength, style='List Bullet')
            
            if business_goal_alignment.get('gaps'):
                doc.add_heading('Gaps to Address', 2)
                for gap in business_goal_alignment['gaps']:
                    doc.add_paragraph(gap, style='List Bullet')
            
            if business_goal_alignment.get('recommendations'):
                doc.add_heading('Strategic Recommendations', 2)
                for i, rec in enumerate(business_goal_alignment['recommendations'], 1):
                    doc.add_paragraph(f"{i}. {rec}")
        
        doc.add_page_break()
        
        # Priority Actions
        priority_actions = ai_rec.get('priority_actions', [])
        if priority_actions:
            doc.add_heading('🚨 Priority Actions', 1)
            for action in priority_actions:
                p = doc.add_paragraph(action)
                if 'CRITICAL' in action:
                    p.runs[0].font.color.rgb = RGBColor(220, 53, 69)
                    p.runs[0].bold = True
                elif 'HIGH' in action:
                    p.runs[0].font.color.rgb = RGBColor(255, 193, 7)
                    p.runs[0].bold = True
        
        # Issues Found
        if issues:
            doc.add_heading(f'⚠️ Issues Found ({len(issues)})', 1)
            for issue in issues:
                doc.add_paragraph(issue, style='List Bullet')
        
        # Suggestions
        if suggestions:
            doc.add_heading(f'✅ Suggested Fixes ({len(suggestions)})', 1)
            for suggestion in suggestions:
                doc.add_paragraph(suggestion, style='List Bullet')
        
        doc.add_page_break()
        
        # Meta Tags Analysis
        meta_tags = metrics.get('meta_tags', {})
        if meta_tags:
            doc.add_heading('📝 Meta Tags Analysis', 1)
            
            doc.add_heading('Title Tag', 2)
            doc.add_paragraph(f"Content: {meta_tags.get('title', 'N/A')}")
            doc.add_paragraph(f"Length: {meta_tags.get('title_length', 0)} characters")
            
            doc.add_heading('Meta Description', 2)
            doc.add_paragraph(f"Content: {meta_tags.get('meta_description', 'N/A')}")
            doc.add_paragraph(f"Length: {meta_tags.get('meta_description_length', 0)} characters")
        
        # Title Suggestions
        title_suggestions = ai_rec.get('title_suggestions', [])
        if title_suggestions:
            doc.add_heading('💡 Title Tag Suggestions', 2)
            for i, suggestion in enumerate(title_suggestions, 1):
                doc.add_paragraph(f"{i}. {suggestion}")
        
        # Meta Description Suggestions
        meta_suggestions = ai_rec.get('meta_description_suggestions', [])
        if meta_suggestions:
            doc.add_heading('💡 Meta Description Suggestions', 2)
            for i, suggestion in enumerate(meta_suggestions, 1):
                doc.add_paragraph(f"{i}. {suggestion}")
        
        doc.add_page_break()
        
        # Keyword Analysis
        keyword_analysis = metrics.get('keyword_analysis', {})
        if keyword_analysis:
            doc.add_heading('🔑 Keyword Analysis', 1)
            
            doc.add_paragraph(f"Primary Keyword: {keyword_analysis.get('primary_keyword', 'N/A')}")
            doc.add_paragraph(f"Keyword Density: {keyword_analysis.get('density', 0)}%")
            doc.add_paragraph(f"In Title: {'✅ Yes' if keyword_analysis.get('in_title') else '❌ No'}")
            doc.add_paragraph(f"In Meta Description: {'✅ Yes' if keyword_analysis.get('in_meta_description') else '❌ No'}")
            doc.add_paragraph(f"In H1: {'✅ Yes' if keyword_analysis.get('in_h1') else '❌ No'}")
            doc.add_paragraph(f"Relevance Score: {keyword_analysis.get('relevance_score', 0)}/100")
            
            top_keywords = keyword_analysis.get('top_keywords', [])
            if top_keywords:
                doc.add_heading('Top Keywords Found', 2)
                doc.add_paragraph(', '.join(top_keywords[:15]))
        
        # Content Analysis
        content = metrics.get('content', {})
        if content:
            doc.add_heading('📄 Content Analysis', 1)
            doc.add_paragraph(f"Word Count: {content.get('word_count', 0)} words")
            doc.add_paragraph(f"Content Quality Score: {content.get('quality_score', 0)}/100")
            doc.add_paragraph(f"Vocabulary Richness: {content.get('vocabulary_richness', 'N/A')}")
            doc.add_paragraph(f"H1 Count: {content.get('h1_count', 0)}")
            doc.add_paragraph(f"H2 Count: {content.get('h2_count', 0)}")
            doc.add_paragraph(f"H3 Count: {content.get('h3_count', 0)}")
        
        # Content Improvements
        content_improvements = ai_rec.get('content_improvements', [])
        if content_improvements:
            doc.add_heading('💡 Content Improvements', 2)
            for i, improvement in enumerate(content_improvements, 1):
                doc.add_paragraph(f"{i}. {improvement}")
        
        # Readability Analysis
        readability = metrics.get('readability', {})
        if readability:
            doc.add_heading('📖 Readability Analysis', 1)
            doc.add_paragraph(f"Readability Score: {readability.get('score', 0)}/100")
            doc.add_paragraph(f"Grade Level: {readability.get('grade_level', 'N/A')}")
        
        # Technical SEO
        technical = metrics.get('technical', {})
        if technical:
            doc.add_heading('⚙️ Technical SEO', 1)
            doc.add_paragraph(f"Page Size: {technical.get('page_size_kb', 0)} KB")
            doc.add_paragraph(f"Response Time: {technical.get('response_time_seconds', 0)}s")
            doc.add_paragraph(f"Mobile Viewport: {'✅ Yes' if technical.get('has_viewport') else '❌ No'}")
            doc.add_paragraph(f"Schema Markup: {'✅ Yes' if technical.get('has_schema') else '❌ No'}")
            doc.add_paragraph(f"Open Graph Tags: {'✅ Yes' if technical.get('has_open_graph') else '❌ No'}")
        
        # Technical Fixes
        technical_fixes = ai_rec.get('technical_fixes', [])
        if technical_fixes:
            doc.add_heading('💡 Technical Fixes', 2)
            for i, fix in enumerate(technical_fixes, 1):
                doc.add_paragraph(f"{i}. {fix}")
        
        # Images
        images = metrics.get('images', {})
        if images:
            doc.add_heading('🖼️ Image Optimization', 1)
            doc.add_paragraph(f"Total Images: {images.get('total', 0)}")
            doc.add_paragraph(f"With Alt Text: {images.get('with_alt', 0)}")
            doc.add_paragraph(f"Without Alt Text: {images.get('without_alt', 0)}")
        
        # Links
        links = metrics.get('links', {})
        if links:
            doc.add_heading('🔗 Link Analysis', 1)
            doc.add_paragraph(f"Internal Links: {links.get('internal', 0)}")
            doc.add_paragraph(f"External Links: {links.get('external', 0)}")
        
        # Competitor Analysis
        competitor_analysis = metrics.get('competitor_analysis')
        if competitor_analysis:
            doc.add_page_break()
            doc.add_heading('🏆 Competitor Keyword Gap Analysis', 1)
            doc.add_paragraph(f"Gap Percentage: {competitor_analysis.get('gap_percentage', 0)}%")
            
            missing = competitor_analysis.get('missing_keywords', [])
            if missing:
                doc.add_heading('Missing Keywords (Target These)', 2)
                doc.add_paragraph(', '.join(missing))
            
            unique = competitor_analysis.get('unique_keywords', [])
            if unique:
                doc.add_heading('Your Unique Keywords', 2)
                doc.add_paragraph(', '.join(unique))
        
        # Competitor Insights
        competitor_insights = ai_rec.get('competitor_insights', [])
        if competitor_insights:
            doc.add_heading('💡 Competitor Insights', 2)
            for i, insight in enumerate(competitor_insights, 1):
                doc.add_paragraph(f"{i}. {insight}")
        
        # Conversion Optimization
        conversion_opt = ai_rec.get('conversion_optimization', [])
        if conversion_opt:
            doc.add_page_break()
            doc.add_heading('💰 Conversion Optimization', 1)
            for i, opt in enumerate(conversion_opt, 1):
                doc.add_paragraph(f"{i}. {opt}")
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def upload_to_drive(self, file_path: str, file_name: Optional[str] = None, folder_id: Optional[str] = None) -> Dict:
        """
        Upload a file to Google Drive.
        
        Args:
            file_path: Local path to file to upload
            file_name: Name for file in Drive (optional, uses local filename)
            folder_id: Google Drive folder ID (optional, uploads to root)
            
        Returns:
            Dictionary with file_id and web_view_link
        """
        if not file_name:
            file_name = os.path.basename(file_path)
        
        file_metadata = {'name': file_name}
        if folder_id:
            file_metadata['parents'] = [folder_id]
        
        media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        file = self.service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink, webContentLink'
        ).execute()
        
        return {
            'file_id': file.get('id'),
            'web_view_link': file.get('webViewLink'),
            'web_content_link': file.get('webContentLink'),
            'file_name': file_name
        }
    
    def create_and_upload_seo_report(
        self, 
        seo_data: Dict, 
        website_name: Optional[str] = None,
        folder_id: Optional[str] = None
    ) -> Dict:
        """
        Create SEO report Word document and upload to Google Drive.
        
        Args:
            seo_data: SEO analysis results
            website_name: Name for the report (e.g., "amazon")
            folder_id: Google Drive folder ID to upload to
            
        Returns:
            Dictionary with upload status and links
        """
        try:
            # Generate filename
            if not website_name:
                url = seo_data.get('url', 'unknown')
                website_name = url.replace('https://', '').replace('http://', '').split('/')[0].split('.')[0]

            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{website_name}_seo_report_{timestamp}.docx"

            # Use platform temp directory to support Windows and *nix
            tmpdir = tempfile.gettempdir()
            local_path = str(Path(tmpdir) / filename)

            # Create Word document
            self._create_word_document(seo_data, local_path)

            # Upload to Google Drive
            upload_result = self.upload_to_drive(local_path, filename, folder_id)

            # Clean up local file
            try:
                if os.path.exists(local_path):
                    os.remove(local_path)
            except Exception:
                # Ignore cleanup failures
                pass

            return {
                'success': True,
                'filename': filename,
                'local_path_created': local_path,
                'file_id': upload_result.get('file_id'),
                'file_url': upload_result.get('web_view_link'),
                'edit_url': upload_result.get('web_view_link'),
                'download_url': upload_result.get('web_content_link'),
                'message': f"SEO report '{filename}' successfully uploaded to Google Drive"
}
        except FileNotFoundError as e:
            # Likely missing credentials.json
            return {
                'success': False,
                'error': str(e),
                'message': (
                    f"Google Drive credentials not found: {e}. "
                    "Place your OAuth2 credentials.json in the project root or provide the path when creating the exporter."
                )
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': f"Failed to create or upload SEO report: {str(e)}"
            }
    
    def list_folders(self) -> list:
        """List all folders in Google Drive"""
        results = self.service.files().list(
            q="mimeType='application/vnd.google-apps.folder' and trashed=false",
            spaces='drive',
            fields='files(id, name)'
        ).execute()
        
        return results.get('files', [])
    
    def create_folder(self, folder_name: str, parent_folder_id: Optional[str] = None) -> str:
        """
        Create a folder in Google Drive.
        
        Args:
            folder_name: Name for the new folder
            parent_folder_id: ID of parent folder (optional)
            
        Returns:
            Folder ID
        """
        file_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        
        if parent_folder_id:
            file_metadata['parents'] = [parent_folder_id]
        
        folder = self.service.files().create(
            body=file_metadata,
            fields='id'
        ).execute()
        
        return folder.get('id')


# Example usage function
def export_seo_to_gdrive(seo_analysis_data: Dict, website_name: str = None) -> Dict:
    """
    Convenience function to export SEO analysis to Google Drive.
    
    Args:
        seo_analysis_data: SEO analysis results from analyze_seo()
        website_name: Name for the website (e.g., "amazon")
        
    Returns:
        Upload result dictionary
    """
    try:
        exporter = GoogleDriveSEOExporter()
        result = exporter.create_and_upload_seo_report(seo_analysis_data, website_name)
        return result
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'message': f"Failed to export to Google Drive: {str(e)}"
        }
